from crewai import Agent, Task, Crew, LLM
from crewai_tools import ScrapeWebsiteTool, SerperDevTool
import os
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
from langchain_community.llms import Ollama

def generate_docx(result):
    """
    Generates a Word document containing AI-generated medical analysis and treatment recommendations.
    """
    doc = Document()
    doc.add_heading("AI-Powered Medical Diagnosis and Treatment Plan", 0)
    doc.add_paragraph(result)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    with open("medical_diagnosis_treatment_plan.docx", "wb") as f:
        f.write(bio.read())
    print("Document saved as medical_diagnosis_treatment_plan.docx")

load_dotenv()
os.environ['SERPER_API_KEY'] = "api key"

search_tool = SerperDevTool()
scrape_tool = ScrapeWebsiteTool()
# Load API key from environment variables
api_key = "your api key"

# Initialize the LLM
llm = LLM(model="gemini/gemini-2.0-flash-exp", api_key=api_key)


gender = input("Select Gender (Male/Female/Other): ")
age = int(input("Enter Age: "))
symptoms = input("Describe Your Symptoms (comma-separated): ")
medical_history = input("Provide Your Medical History (comma-separated): ")

diagnostician = Agent(
    role="Medical Diagnostician",
    goal="Analyze patient symptoms and medical history to provide a preliminary diagnosis.",
    backstory="This AI agent specializes in diagnosing medical conditions based on reported symptoms and history.",
    verbose=True,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm
)

treatment_advisor = Agent(
    role="Treatment Advisor",
    goal="Recommend appropriate treatment plans based on the diagnosis provided by the Medical Diagnostician.",
    backstory="This AI agent tailors treatment plans based on diagnosis, history, and best medical practices.",
    verbose=True,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm
)

diagnose_task = Task(
    description=f"""
        1. Analyze the patient's symptoms ({symptoms}) and medical history ({medical_history}).
        2. Provide a preliminary diagnosis with possible conditions.
        3. Limit the diagnosis to the most likely conditions.
    """,
    expected_output="A preliminary diagnosis with a list of possible conditions.",
    agent=diagnostician
)

treatment_task = Task(
    description=f"""
        1. Based on the diagnosis, recommend appropriate treatment plans step by step.
        2. Consider the patient's medical history ({medical_history}) and current symptoms ({symptoms}).
        3. Provide detailed treatment recommendations, including medications, lifestyle changes, and follow-up care.
    """,
    expected_output="A comprehensive treatment plan tailored to the patient's needs.",
    agent=treatment_advisor
)

crew = Crew(
    agents=[diagnostician, treatment_advisor],
    tasks=[diagnose_task, treatment_task],
    verbose=2
)

print("Generating AI-Powered Diagnosis and Treatment Plan...")
results = crew.kickoff(inputs={"symptoms": symptoms, "medical_history": medical_history})
print(results)
generate_docx(results)
